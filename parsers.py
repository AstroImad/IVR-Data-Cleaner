"""
Parses IVR script documents (PDF/DOCX) to extract:
- Questions (mapped to flow numbers)
- Answer choices (mapped to FlowNo_X=Y patterns)
- Flow graph (redirect relationships between flows for skip logic handling)

Handles multi-layer/branching IVR scripts where:
- "Tekan X untuk Y Call flow N" indicates an answer that redirects to another flow
- "Tekan N hingga M" is a range description for multi-item sub-questions
- Duplicate question texts are disambiguated by appending the flow number
- Likert scale matrices (e.g. Bomba, Polis) maintain their overarching context
"""

import re
from typing import Any, Dict, Tuple, Optional, List, Set


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    import pdfplumber
    import io
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract paragraph and table text from DOCX bytes.

    IVR scripts are commonly authored as paragraphs, but routing matrices and
    answer tables are also valid Word layouts.  Include table rows so the
    parser sees the same routing vocabulary regardless of the Word layout.
    """
    from docx import Document
    import io
    doc = Document(io.BytesIO(file_bytes))
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.replace("\n", " ") for cell in row.cells))
    return "\n".join(parts)


def clean_flow_line(text: str) -> str:
    """Normalize flow labels and common navigation separators."""
    text = text.replace("\u2192", " -> ").replace("\u2794", " -> ")
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r'Call\s*\n?\s*flow\s*\n?\s*(\d+)', r'Call flow \1', text, flags=re.IGNORECASE)
    text = re.sub(r'Call\s*flow\s*(\d+)', r'Call flow \1', text, flags=re.IGNORECASE)
    text = re.sub(r'\bcallflow\s*(\d+)', r'Call flow \1', text, flags=re.IGNORECASE)
    text = re.sub(r'\bflow\s*#?\s*(\d+)\b', r'flow \1', text, flags=re.IGNORECASE)
    return text


def _disambiguate_questions(flow_to_question: Dict[int, str]) -> Dict[int, str]:
    """Ensure all question texts are unique by appending flow number to duplicates.

    Example:
        Input:  {5: "Berpuas hati?", 6: "Berpuas hati?", 7: "Umur anda?"}
        Output: {5: "Berpuas hati? (Call flow 5)", 6: "Berpuas hati? (Call flow 6)", 7: "Umur anda?"}"""

    from collections import Counter

    # Count how many flows share each question text
    question_counts = Counter(flow_to_question.values())

    # Collect question texts that appear more than once
    duplicates_questions = {q for q, c in question_counts.items() if c > 1}

    # If no duplicate exist, return unchanged
    if not duplicates_questions:
        return flow_to_question

    # Build a new dict, appending flow number only to duplicate questions
    result = {}
    for flow_num, question in flow_to_question.items():
        if question in duplicates_questions:
            result[flow_num] = f"{question} (Call flow {flow_num})"
        else:
            result[flow_num] = question
    return result


def _build_flow_graph(
    original_text: str,
    processed_text: str,
    cf_matches: list,
    tekan_untuk_pattern_loose,
    tekan_range_pattern,
) -> Dict[int, Dict]:
    """Build a normalized flow graph from several IVR routing notations.

    A route is represented as ``{"choice": int | None, "target": int | None,
    "label": str, "terminal": bool, "source": str}``.  The legacy
    ``answer_redirects`` mapping is populated at the same time for callers
    that already use the original parser contract.
    """
    standalone_flow_nums = {int(m.group(1)) for m in cf_matches}

    lines = original_text.split('\n')
    current_flow = None
    flow_graph: Dict[int, Dict] = {}
    all_redirect_targets: Set[int] = set()

    standalone_cf_re = re.compile(r'\bCall\s+flow\s+(\d+)\b', re.IGNORECASE)
    flow_target_re = re.compile(
        r'\b(?:call\s+)?flow\s*#?\s*(\d+)\b', re.IGNORECASE
    )
    choice_re = re.compile(
        r'\b(?:tekan|press|pilihan|option|jawapan|answer)\s*[:=]?\s*(\d+)\b',
        re.IGNORECASE,
    )
    terminal_re = re.compile(
        r'\b(?:tamat|terima\s+kasih|end(?:\s+of\s+survey)?|terminate|hang\s*up|stop)\b',
        re.IGNORECASE,
    )
    navigation_re = re.compile(
        r'(?:teruskan\s+ke|pergi\s+ke|lompat\s+ke|goto|go\s+to|skip\s+to|next\s+to|route\s+to)\s*',
        re.IGNORECASE,
    )
    range_choice_re = re.compile(
        r'\b(?:tekan|press|pilihan|option)\s+\d+\s*(?:hingga|to|-)\s*\d+\b',
        re.IGNORECASE,
    )

    def ensure_flow(flow_num: int) -> Dict[str, Any]:
        return flow_graph.setdefault(flow_num, {
            'answer_redirects': {},
            'routes': [],
            'terminal_choices': [],
            'terminal': False,
            'is_answer_branch': False,
        })

    def add_route(flow_num: int, choice: Optional[int], target: Optional[int],
                  label: str, source: str, terminal: bool = False) -> None:
        info = ensure_flow(flow_num)
        route = {
            'choice': choice,
            'target': target,
            'label': label.strip(' .:-'),
            'terminal': terminal,
            'source': source.strip(),
        }
        if route not in info['routes']:
            info['routes'].append(route)
        if choice is not None and target is not None:
            info['answer_redirects'][choice] = target
            all_redirect_targets.add(target)
        if terminal:
            if choice is not None and choice not in info['terminal_choices']:
                info['terminal_choices'].append(choice)
            info['terminal'] = True

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # A flow marker establishes the context for subsequent answer lines.
        choice_match = choice_re.search(line_stripped)
        choice = int(choice_match.group(1)) if choice_match else None
        target_matches = list(flow_target_re.finditer(line_stripped))
        # The first flow number can be the current-flow label; use the last
        # explicit target when a line contains both source and destination.
        target = int(target_matches[-1].group(1)) if target_matches else None

        cf_matches_on_line = list(standalone_cf_re.finditer(line_stripped))
        if cf_matches_on_line and not (choice is not None and target is not None):
            marker_flow = int(cf_matches_on_line[-1].group(1))
            if marker_flow in standalone_flow_nums:
                current_flow = marker_flow
                ensure_flow(current_flow)

        if current_flow is None:
            continue

        has_navigation = bool(navigation_re.search(line_stripped))
        has_arrow = '->' in line_stripped
        is_terminal = bool(terminal_re.search(line_stripped)) and target is None

        # A line such as "Tekan 1 hingga 3 Call flow 5" describes a matrix
        # response range, not a single routing choice.
        is_range_declaration = bool(range_choice_re.search(line_stripped))
        is_self_route = target == current_flow and choice is None

        if target is not None and not is_range_declaration and not is_self_route and (
            choice is not None or has_navigation or has_arrow
        ):
            label = line_stripped
            if choice_match:
                label = line_stripped[choice_match.end():]
            label = re.sub(flow_target_re, '', label).strip()
            add_route(current_flow, choice, target, label, line_stripped)
        elif is_terminal and (choice is not None or line_stripped.lower().startswith(('tamat', 'terima'))):
            label = line_stripped[choice_match.end():] if choice_match else line_stripped
            add_route(current_flow, choice, None, label, line_stripped, terminal=True)

    for target_flow in all_redirect_targets:
        if target_flow in flow_graph:
            flow_graph[target_flow]['is_answer_branch'] = True

    return flow_graph


def _get_core_question_text(text: str) -> str:
    """Strip 'Soalan [ordinal].' prefix from question text."""
    stripped = re.sub(r'^Soalan\s+\w+(\s+\w+)*\.\s*', '', text, flags=re.IGNORECASE)
    return stripped.strip() if stripped.strip() else text.strip()


def _identify_branch_groups(
    flow_graph: Dict[int, Dict],
    flow_to_question: Dict[int, str],
    flow_value_mapping: Dict[str, str],
) -> List[List[int]]:
    """Identify groups of mutually exclusive branch flows."""
    core_to_flows: Dict[str, List[int]] = {}
    for flow_num, question in flow_to_question.items():
        core = _get_core_question_text(question)
        if core not in core_to_flows:
            core_to_flows[core] = []
        core_to_flows[core].append(flow_num)

    branch_groups: List[List[int]] = []
    for core, flows in core_to_flows.items():
        if len(flows) >= 2:
            branch_groups.append(sorted(flows))

    parent_redirect_groups: Dict[int, List[int]] = {}
    for flow_num, info in flow_graph.items():
        redirects = info.get('answer_redirects', {})
        if len(redirects) >= 2:
            targets = list(set(redirects.values()))
            if len(targets) >= 2:
                parent_redirect_groups[flow_num] = targets

    for parent, targets in parent_redirect_groups.items():
        sorted_targets = sorted(targets)
        merged = False
        for existing in branch_groups:
            if len(set(existing) & set(sorted_targets)) >= 2:
                existing_set = set(existing) | set(sorted_targets)
                branch_groups[branch_groups.index(existing)] = sorted(existing_set)
                merged = True
                break
        if not merged:
            branch_groups.append(sorted_targets)

    new_groups: List[List[int]] = []
    for group in branch_groups:
        merge_targets: Dict[int, List[int]] = {}  
        for flow_num in group:
            if flow_num not in flow_graph:
                continue
            redirects = flow_graph[flow_num].get('answer_redirects', {})
            if not redirects:
                continue
            unique_targets = set(redirects.values())
            if len(unique_targets) == 1:
                target = list(unique_targets)[0]
                if target not in merge_targets:
                    merge_targets[target] = []
                merge_targets[target].append(flow_num)

        if len(merge_targets) >= 2:
            downstream_group = sorted(merge_targets.keys())
            new_groups.append(downstream_group)

    for new_group in new_groups:
        merged = False
        for existing in branch_groups:
            if len(set(existing) & set(new_group)) >= 2:
                existing_set = set(existing) | set(new_group)
                branch_groups[branch_groups.index(existing)] = sorted(existing_set)
                merged = True
                break
        if not merged:
            branch_groups.append(new_group)

    unique_groups = []
    seen = set()
    for group in branch_groups:
        key = tuple(sorted(group))
        if key not in seen:
            seen.add(key)
            unique_groups.append(sorted(group))

    return unique_groups


def _extract_answer_mappings_from_lines(text: str) -> Dict[str, str]:
    """Extract answer codes using line-level flow context.

    Splitting a document only at ``Call flow`` markers is ambiguous because
    routing destinations also contain that marker.  A line-oriented pass
    avoids assigning the next matrix row to the previous flow and preserves
    the source flow for ordinary redirect answers.
    """
    marker_re = re.compile(r'\bCall\s+flow\s+(\d+)\b', re.IGNORECASE)
    choice_re = re.compile(
        r'^\s*(?:Tekan|Press|Pilihan|Option|Jawapan|Answer)\s+'
        r'(\d+)\s+(?:(?:untuk|for|to)\s+)?(.+)$',
        re.IGNORECASE,
    )
    range_re = re.compile(
        r'^(.*?)\b(?:tekan|press|pilihan|option)\s+(\d+)\s*'
        r'(?:hingga|to|-)\s*(\d+)\b',
        re.IGNORECASE,
    )
    mapping: Dict[str, str] = {}
    current_flow: Optional[int] = None

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        markers = list(marker_re.finditer(stripped))
        target = int(markers[-1].group(1)) if markers else None
        range_match = range_re.search(stripped)
        choice_match = choice_re.match(stripped)

        if range_match and target is not None:
            label = range_match.group(1).strip(' .:-\t')
            for choice in range(int(range_match.group(2)), int(range_match.group(3)) + 1):
                mapping.setdefault(f"FlowNo_{target}={choice}", label)
            current_flow = target
            continue

        if choice_match and current_flow is not None:
            choice = int(choice_match.group(1))
            answer = choice_match.group(2).strip()
            answer = re.sub(marker_re, '', answer).strip(' .:-')
            mapping[f"FlowNo_{current_flow}={choice}"] = answer

        # A question/flow declaration starts the next flow.  Answer lines
        # that point elsewhere deliberately leave the source flow active.
        if target is not None and not choice_match:
            current_flow = target

    return mapping


def parse_ivr_script(file_bytes: bytes, filename: str) -> Tuple[
    Dict[int, str],
    Dict[str, str],
    Dict[int, Dict],
    List[List[int]],
]:
    """
    Parse an IVR script document to extract questions, answer mappings,
    flow graph, and branch groups.
    """
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}")

    # Fix multi-line and compact flow labels before any parsing pass.
    text = clean_flow_line(text)

    # ââ Regex patterns âââââââââââââââââââââââââââââââââââââââââââââââââââ
    # STRICT PATTERNS: Use ^\s* to ensure it only matches at the start of a line.
    tekan_redirect_pattern = re.compile(
        r'^\s*(Tekan\s+\d+\s+untuk\s+.+?)\s+Call\s+flow\s+\d+',
        re.IGNORECASE | re.MULTILINE
    )

    tekan_untuk_pattern_strict = re.compile(
        r'^\s*(?:Tekan|Press|Pilihan|Option|Jawapan|Answer)\s+(\d+)\s+(?:untuk|for|to)?\s*(.+)',
        re.IGNORECASE
    )

    tekan_untuk_pattern_loose = re.compile(
        r'(?:Tekan|Press|Pilihan|Option|Jawapan|Answer)\s+(\d+)\s+(?:untuk|for|to)?\s*(.+)',
        re.IGNORECASE
    )

    # likert questions pattern for answer
    tekan_range_pattern = re.compile(r'Tekan\s+\d+\s+hingga\s+\d+', re.IGNORECASE)

    # liker questions pattern
    multi_item_pattern = re.compile(
        r'(.+?)\s+tekan\s+(\d+)\s+hingga\s+(\d+)\s+Call\s+flow\s+(\d+)',
        re.IGNORECASE
    )

    skip_patterns = [
        'salam sejahtera', 'terima kasih', 'kajian bebas', 'cpi',
        'hanya merangkumi', 'soalan untuk bukan pengundi',
        'berdasarkan', 'jawab soalan', 'jawab ini',
    ]

    def replace_redirect(match):
        """Replace 'Call flow M' in Tekan lines with a placeholder."""
        tekan_text = match.group(1) 
        tekan_match = tekan_untuk_pattern_loose.search(tekan_text)
        if tekan_match:
            return tekan_text 
        return match.group(0)

    processed_text = tekan_redirect_pattern.sub(replace_redirect, text)

    # ââ First pass: detect multi-item sub-questions ââââââââââââââââââââââ
    multi_item_questions: Dict[int, str] = {}

    for match in multi_item_pattern.finditer(processed_text):
        entity_name = match.group(1).strip()
        flow_num = int(match.group(4))
        multi_item_questions[flow_num] = entity_name

    # ââ Second pass: split by standalone "Call flow N" and process ââââââ
    call_flow_pattern = re.compile(r'\bCall\s+flow\s+(\d+)\b', re.IGNORECASE)
    cf_matches = list(call_flow_pattern.finditer(processed_text))

    flow_to_question: Dict[int, str] = {}
    flow_value_mapping: Dict[str, str] = {}

    flow_graph = _build_flow_graph(
        text, processed_text, cf_matches, tekan_untuk_pattern_loose, tekan_range_pattern
    )

    last_seen_q = ""

    for idx, cf_match in enumerate(cf_matches):
        flow_num = int(cf_match.group(1))
        cf_end = cf_match.end()

        if idx > 0:
            content_before = processed_text[cf_matches[idx - 1].end():cf_match.start()]
        else:
            content_before = processed_text[:cf_match.start()]

        if idx + 1 < len(cf_matches):
            content_after = processed_text[cf_end:cf_matches[idx + 1].start()]
        else:
            content_after = processed_text[cf_end:]

        # ââ Extract question text ââââââââââââââââââââââââââââââââââââââ
        q = _extract_question_from_content(
            content_before, content_after,
            tekan_untuk_pattern_strict, tekan_range_pattern, skip_patterns
        )

        # Context pairing for Likert matrices (e.g. Bomba, Polis)
        if flow_num in multi_item_questions:
            if q:
                last_seen_q = q
            if last_seen_q:
                flow_to_question[flow_num] = f"{last_seen_q} [{multi_item_questions[flow_num]}]"
            else:
                flow_to_question[flow_num] = multi_item_questions[flow_num]
        else:
            if q:
                last_seen_q = q
            if q and flow_num >= 2:
                flow_to_question[flow_num] = q

        # ââ Extract answer mappings from content_after âââââââââââââââââ
        _extract_answers_from_content(
            content_after, flow_num,
            tekan_untuk_pattern_strict, tekan_range_pattern,
            flow_value_mapping
        )

    # Rebuild mappings with the unambiguous line-oriented pass.  The legacy
    # content-window extraction above remains useful for unusual paragraph
    # layouts, while this pass prevents repeated matrix flows from overwriting
    # an earlier row's codebook.
    flow_value_mapping.update(_extract_answer_mappings_from_lines(text))

    branch_groups = _identify_branch_groups(flow_graph, flow_to_question, flow_value_mapping)
    flow_to_question = _disambiguate_questions(flow_to_question)

    # A terminal flow may be reached by a redirect but have no answer lines.
    # Mark it from its question/script text so downstream code can traverse
    # routes without relying on a hard-coded flow number.
    terminal_text_re = re.compile(
        r'\b(?:terima\s+kasih|tamat|end(?:\s+of\s+survey)?|terminate|hang\s*up)\b',
        re.IGNORECASE,
    )
    marker_lines_by_flow: Dict[int, List[str]] = {}
    for line in text.split('\n'):
        marker_matches = list(re.finditer(r'\bCall\s+flow\s+(\d+)\b', line, re.IGNORECASE))
        for marker_match in marker_matches:
            marker_lines_by_flow.setdefault(int(marker_match.group(1)), []).append(line)

    for flow_num, question in flow_to_question.items():
        section_text = ' '.join(marker_lines_by_flow.get(flow_num, []))
        if terminal_text_re.search(question) or terminal_text_re.search(section_text):
            flow_graph.setdefault(flow_num, {
                'answer_redirects': {}, 'routes': [],
                'terminal_choices': [], 'terminal': False,
                'is_answer_branch': False,
            })['terminal'] = True

    return flow_to_question, flow_value_mapping, flow_graph, branch_groups


def get_skip_logic_candidates(
    flow_graph: Dict[int, Dict],
    flow_to_question: Optional[Dict[int, str]] = None,
) -> List[Dict[str, Any]]:
    """Return route choices that are useful as skip/branch candidates.

    This intentionally reports routes instead of silently filtering data:
    routing to another flow is not necessarily a disqualifying skip.  The UI
    can therefore present every detected candidate, including terminal routes,
    alternate/non-voter paths, and ordinary branch routes for manual review.
    ``terminal_reached`` follows known route edges until it reaches a terminal
    flow, while ``alternate`` highlights common alternate-path wording.
    """
    flow_to_question = flow_to_question or {}
    terminal_words = re.compile(
        r'\b(?:tamat|terima\s+kasih|end(?:\s+of\s+survey)?|terminate|hang\s*up)\b',
        re.IGNORECASE,
    )
    alternate_words = re.compile(
        r'\b(?:tambahan|bukan\s+pengundi|non[- ]?voter|alternate|screen|saringan|skip)\b',
        re.IGNORECASE,
    )
    screening_question = re.compile(
        r'\b(?:pengundi|voter|berdaftar|registered)\b',
        re.IGNORECASE,
    )
    negative_answer = re.compile(
        r'\b(?:tidak|tak|no|bukan|not)\b',
        re.IGNORECASE,
    )

    def reaches_terminal(flow_num: Optional[int], visited: Set[int]) -> bool:
        if flow_num is None or flow_num in visited:
            return False
        visited.add(flow_num)
        info = flow_graph.get(flow_num, {})
        if info.get('terminal') or terminal_words.search(flow_to_question.get(flow_num, '')):
            return True
        return any(
            reaches_terminal(route.get('target'), visited.copy())
            for route in info.get('routes', [])
            if route.get('target') is not None
        )

    candidates: List[Dict[str, Any]] = []
    for source_flow, info in sorted(flow_graph.items()):
        source_text = flow_to_question.get(source_flow, '')
        for route in info.get('routes', []):
            target = route.get('target')
            target_text = flow_to_question.get(target, '') if target is not None else ''
            alternate = bool(
                alternate_words.search(route.get('label', ''))
                or alternate_words.search(source_text)
                or alternate_words.search(target_text)
                or (
                    screening_question.search(source_text)
                    and negative_answer.search(route.get('label', ''))
                )
            )
            candidates.append({
                'flow': source_flow,
                'choice': route.get('choice'),
                'value': (
                    f"FlowNo_{source_flow}={route['choice']}"
                    if route.get('choice') is not None else None
                ),
                'target': target,
                'label': route.get('label', ''),
                'source': route.get('source', ''),
                'terminal': bool(route.get('terminal')),
                'terminal_reached': reaches_terminal(target, set()),
                'alternate': alternate,
            })
    return candidates


def _extract_question_from_content(
    content_before: str,
    content_after: str,
    tekan_untuk_pattern_strict,
    tekan_range_pattern,
    skip_patterns: List[str]
) -> str:
    """Extract question text from content before and after a 'Call flow N' marker."""
    before_lines = content_before.strip().split('\n')
    question_parts = []

    for line in before_lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        line_lower = line_stripped.lower()
        if any(skip in line_lower for skip in skip_patterns):
            continue
        if re.match(r'^Q\d+', line_stripped, re.IGNORECASE):
            continue
        if tekan_untuk_pattern_strict.search(line_stripped):
            continue
        if tekan_range_pattern.search(line_stripped):
            continue

        line_stripped = re.sub(r'^[\]\[\)\(}{\s]+', '', line_stripped).strip()
        if not line_stripped:
            continue

        question_parts.append(line_stripped)

    question_text = " ".join(question_parts).strip()

    after_lines = content_after.strip().split('\n')
    trailing_question = []

    for line in after_lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        if tekan_untuk_pattern_strict.search(line_stripped):
            break
        if tekan_range_pattern.search(line_stripped):
            break
            
        line_lower = line_stripped.lower()
        if any(skip in line_lower for skip in skip_patterns):
            continue
        if re.match(r'^Q\d+', line_stripped, re.IGNORECASE):
            continue
            
        trailing_question.append(line_stripped)

    if trailing_question:
        trailing_text = " ".join(trailing_question).strip()
        if question_text:
            question_text = question_text.rstrip() + " " + trailing_text
        else:
            question_text = trailing_text

    return question_text


def _extract_answers_from_content(
    content_after: str,
    flow_num: int,
    tekan_untuk_pattern_strict,
    tekan_range_pattern,
    flow_value_mapping: Dict[str, str]
):
    """Extract answer mappings from content after a 'Call flow N' marker."""
    after_lines = content_after.strip().split('\n')

    range_pattern = re.compile(
        r'(?P<label>.*?)\b(?:tekan|press|pilihan|option)\s+(?P<start>\d+)\s*(?:hingga|to|-)\s*(?P<end>\d+)\b',
        re.IGNORECASE,
    )

    for line in after_lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        range_match = range_pattern.search(line_stripped)
        if range_match:
            label = range_match.group('label').strip(' .:-\t')
            for choice_num in range(
                int(range_match.group('start')),
                int(range_match.group('end')) + 1,
            ):
                key = f"FlowNo_{flow_num}={choice_num}"
                if label:
                    flow_value_mapping[key] = label
            continue

        tekan_match = tekan_untuk_pattern_strict.search(line_stripped)
        if tekan_match:
            choice_num = int(tekan_match.group(1))
            answer_text = tekan_match.group(2).strip()
            answer_text = re.sub(r'\s*Call\s+flow\s+\d+\s*$', '', answer_text, flags=re.IGNORECASE).strip()
            key = f"FlowNo_{flow_num}={choice_num}"
            flow_value_mapping[key] = answer_text